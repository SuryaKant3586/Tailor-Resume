import os
import docx
import tempfile
import pypandoc

import logging
logging.basicConfig(
    format='%(asctime)s - %(message)s',
    level=logging.INFO,
    handlers=[
        logging.FileHandler("app.log"),
        logging.StreamHandler()
    ]
)

def replace_placeholders(doc, replacements):
    def replace_in_text(text, replacements):
        for key, value in replacements.items():
            text = text.replace(key, value)
        return text

    for paragraph in doc.paragraphs:
        # Concatenate all runs to form the complete paragraph text
        full_text = ''.join([run.text for run in paragraph.runs])
        new_text = replace_in_text(full_text, replacements)

        # Assign new text back to runs
        if new_text != full_text:
            paragraph.clear()  # Clear the current paragraph
            paragraph.add_run(new_text)  # Add the new, updated text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders(cell, replacements)  # Recursive call for table cells


def generate_coverletter_files(coverletter_json: dict, resume_json: dict):
    logging.info("generate_coverletter_files() called.")
    doc_file = None
    pdf_file = None  # Initialize with None to avoid unassigned variable issues

    coverletter_template_path = "engine/utils/templates/coverletter_template.docx"
    # Check if the template file exists
    if not os.path.exists(coverletter_template_path):
        logging.error(f"Error: Cover letter template file not found at {coverletter_template_path}")
        return None, None  # Return None for both files if template is missing

    try:
        # Safely access the JSON keys with default empty strings if they are missing
        replacements = {
            "{{name}}": resume_json.get('basics', {}).get("name", ""),
            "{{location}}": resume_json.get('basics', {}).get("address", ""),
            "{{email}}": resume_json.get('basics', {}).get("email", ""),
            "{{phone}}": resume_json.get('basics', {}).get("phone", ""),
            "{{website}}": resume_json.get('basics', {}).get("website", ""),
            "{{coverletter_content}}": coverletter_json.get("coverletter", "")
        }
        
        # Load the Word document template
        doc = docx.Document(coverletter_template_path)
        
        # Replace placeholders with actual values
        replace_placeholders(doc, replacements)

        with tempfile.TemporaryDirectory() as temp_dir:
            doc_path = os.path.join(temp_dir, "coverletter.docx")
            doc.save(doc_path)
            with open(doc_path, "rb") as f:
                doc_file = f.read()
            logging.debug("- tailored cover letter DOCX file generated -")

            pdf_path = os.path.join(temp_dir, "coverletter.pdf")
            pypandoc.convert_file(doc_path, 'pdf', outputfile=pdf_path)
            with open(pdf_path, "rb") as f:
                pdf_file = f.read()
            logging.debug("- tailored cover letter PDF file generated -")

    except Exception as e:
        logging.error(f"Error: error generating cover letter files.\nDetailed error: {e}")

    return (doc_file, 
            pdf_file)
