import tempfile
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pypandoc


import logging
logging.basicConfig(
    format='%(asctime)s - %(message)s',
    level=logging.INFO,
    handlers=[
        logging.FileHandler("app.log"),
        logging.StreamHandler()
    ]
)

def add_divider(doc, space_before=Pt(0), space_after=Pt(0)):
    # Add an empty paragraph to act as the divider
    paragraph = doc.add_paragraph()
    p_pr = paragraph._element.get_or_add_pPr()

    # Create and configure the bottom border element
    p_bdr = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '4')  # Border size
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), '000000')  # Border color

    # Append the bottom border to the paragraph properties
    p_bdr.append(bottom_border)
    p_pr.append(p_bdr)

    # Adjust spacing around the divider line
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = space_before
    paragraph_format.space_after = space_after


def gen_resume_docx(resume_json: dict):
    # Create a DOCX Document
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # For A4 page size
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Name and Contact Information
    name_paragraph = doc.add_paragraph(resume_json["basics"]["name"])
    name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_paragraph.runs[0].font.size = Pt(20)
    name_paragraph.runs[0].bold = True

    contact_info = f"{resume_json['basics']['email']} | {resume_json['basics']['phone']} | {resume_json['basics']['address']}"
    contact_paragraph = doc.add_paragraph(contact_info)
    contact_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_paragraph.space_after = Pt(12)

    # Summary Section
    summary_paragraph = doc.add_paragraph(resume_json["summary"]["text"])
    summary_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    summary_paragraph.space_after = Pt(6)

    add_divider(doc)

    # Education Section
    doc.add_heading("Education", level=1)
    for edu in resume_json["education"]:
        edu_heading = doc.add_heading(f"{edu['studyType']} in {edu['area']}", level=2)
        edu_heading.runs[0].font.size = Pt(12)
        edu_heading.runs[0].bold = True

        edu_details = doc.add_paragraph()
        edu_details.add_run(f"{edu['institution']}, {edu['date']}").italic = True
        edu_details.space_after = Pt(3)

    add_divider(doc)

    # Work Experience Section
    doc.add_heading("Work Experience", level=1)
    for job in resume_json["work"]:
        job_heading = doc.add_heading(f"{job['company']} - {job['position']}", level=2)
        job_heading.runs[0].font.size = Pt(12)
        job_heading.runs[0].bold = True

        dates_paragraph = doc.add_paragraph(f"{job['startDate']} - {job['endDate']}")
        dates_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        dates_paragraph.italic = True

        for highlight in job["highlights"]:
            bullet = doc.add_paragraph(f"{highlight}", style="List Bullet")
            bullet.paragraph_format.left_indent = Inches(0.25)
        dates_paragraph.space_after = Pt(3)

    add_divider(doc)

    # Projects Section
    doc.add_heading("Projects", level=1)
    for project in resume_json["projects"]:
        project_heading = doc.add_heading(project["name"], level=2)
        project_heading.runs[0].font.size = Pt(12)
        project_heading.runs[0].bold = True

        for desc in project["description"]:
            bullet = doc.add_paragraph(f"{desc}", style="List Bullet")
            bullet.paragraph_format.left_indent = Inches(0.25)

    add_divider(doc)

    # Skills Section
    doc.add_heading("Skills", level=1)
    for skill in resume_json["skills"]:
        if skill.get("keywords"):  # Only include skills with keywords
            skill_entry = doc.add_paragraph()
            skill_entry.add_run(f"{skill['name']}: ").bold = True
            skill_entry.add_run(", ".join(skill["keywords"]))
            skill_entry.space_after = Pt(3)

    return doc


def generate_resume_files(resume_json: dict):
    logging.info("generate_resume_files() called.")
    doc_file = None
    pdf_file = None  # Initialize with None to avoid unassigned variable issues

    try:
        # Download Pandoc if not already installed
        pypandoc.download_pandoc()

        doc = gen_resume_docx(resume_json)

        with tempfile.TemporaryDirectory() as temp_dir:
            doc_path = f"{temp_dir}/resume.docx"
            doc.save(doc_path)
            with open(doc_path, "rb") as f:
                doc_file = f.read()
            logging.debug("- tailored resume DOCX file generated -")

            pdf_path = f"{temp_dir}/resume.pdf"
            pypandoc.convert_file(doc_path, 'pdf', outputfile=pdf_path)
            with open(pdf_path, "rb") as f:
                pdf_file = f.read()
            logging.debug("- tailored resume PDF file generated -")

    except Exception as e:
        logging.error(f"Error: error generating tailored resume files.\nDetailed error: {e}")

    return (doc_file,
            pdf_file)
